from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Pharmacovigilance & Regulatory Intelligence System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Technical Documentation & KT Guide').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Generated automatically for developer onboarding.\n')
    
    # Section 1: Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Welcome to the AI Pharmacovigilance (PV) project. This document serves as a "
        "low-level, step-by-step Knowledge Transfer (KT) guide. It explains the inner "
        "workings of the application, how the frontend and backend communicate, "
        "and exactly what happens during the data processing pipeline."
    )
    
    # Section 2: Architecture & Tech Stack
    doc.add_heading('2. Architecture & Tech Stack', level=1)
    doc.add_paragraph("The application follows a decoupled client-server architecture:")
    
    p = doc.add_paragraph()
    p.add_run("Frontend: ").bold = True
    p.add_run("React 19, Vite, Tailwind CSS, Recharts for analytics, React Router DOM.")
    
    p = doc.add_paragraph()
    p.add_run("Backend: ").bold = True
    p.add_run("FastAPI (Python 3.x), SciSpacy (Medical NLP), Supabase (PostgreSQL), Milvus Lite (Vector DB), OpenFDA API, Integrated LLM Service, Langfuse (Observability).")
    
    # Section 3: Directory Structure
    doc.add_heading('3. Key Directory Structure', level=1)
    doc.add_heading('Backend Structure (/backend)', level=2)
    doc.add_paragraph("app/routes: Contains all the FastAPI endpoints (analyze.py, report.py, fda.py, etc.)", style='List Bullet')
    doc.add_paragraph("app/services: Contains the core business logic (llm_service.py, scispacy_service.py, fda_service.py, etc.)", style='List Bullet')
    doc.add_paragraph("milvus/: Contains the Milvus vector database local setup/data.", style='List Bullet')
    doc.add_paragraph("main.py: The entry point of the FastAPI application.", style='List Bullet')

    doc.add_heading('Frontend Structure (/frontend_new)', level=2)
    doc.add_paragraph("src/components: Reusable UI widgets and elements.", style='List Bullet')
    doc.add_paragraph("src/pages: Page-level components mapping to specific routes.", style='List Bullet')
    doc.add_paragraph("src/layouts: Layout components like DashboardLayout.jsx.", style='List Bullet')

    # Section 4: Low-Level Step-by-Step Workflow
    doc.add_heading('4. Low-Level Step-by-Step Data Pipeline', level=1)
    
    doc.add_heading('Step 4.1: Document Upload & Case Intake', level=2)
    doc.add_paragraph("Endpoint: ", style='List Bullet').add_run("POST /analyze/").bold = True
    doc.add_paragraph(
        "Workflow: When a user uploads a document or pastes text in the frontend, it hits this endpoint. "
        "The backend reads the raw content. If it's a multi-patient document, the segmentation service (app/services/segmentation/) slices it into individual cases."
    )
    
    doc.add_heading('Step 4.2: Medical Entity Extraction (NLP)', level=2)
    doc.add_paragraph("Service: ", style='List Bullet').add_run("scispacy_service.py").bold = True
    doc.add_paragraph(
        "Workflow: The individual case text is passed to the SciSpacy NLP model. "
        "It extracts raw terms. These terms are categorized into Drugs, Symptoms, and Conditions. "
        "A secondary layer, the drug_validator_service.py, cleans these extracted terms to ensure accuracy."
    )

    doc.add_heading('Step 4.3: Knowledge Base & Vectorization (RAG)', level=2)
    doc.add_paragraph("Service: ", style='List Bullet').add_run("rag_service.py, embedding_service.py").bold = True
    doc.add_paragraph(
        "Workflow: A unique UUID (analysis_id) is generated. The case text and extracted entities are converted into vector embeddings. "
        "These are stored in the Milvus Vector DB (input_documents collection). The relational data is saved to Supabase."
    )

    doc.add_heading('Step 4.4: Report Generation & Suspect Drug Identification', level=2)
    doc.add_paragraph("Endpoint: ", style='List Bullet').add_run("POST /report/generate").bold = True
    doc.add_paragraph("Service: ", style='List Bullet').add_run("llm_service.py, fda_service.py").bold = True
    doc.add_paragraph(
        "Workflow: The user triggers the report generation. The LLM reads the narrative to identify the 'Primary Suspect Drug'. "
        "Once identified, fda_service.py queries the OpenFDA API with the drug and symptoms to find historical signals and related data."
    )

    doc.add_heading('Step 4.5: LLM Reasoning & Assessment', level=2)
    doc.add_paragraph("Service: ", style='List Bullet').add_run("llm_service.py").bold = True
    doc.add_paragraph(
        "Workflow: A massive context prompt is sent to the LLM containing: the text, SciSpacy entities, FDA data, and RAG context. "
        "The LLM performs Causality Assessment (did the drug cause this?) and Seriousness Assessment (is it life-threatening?). "
        "It also formulates a chronological timeline."
    )

    doc.add_heading('Step 4.6: Evidence Mapping & Formatting', level=2)
    doc.add_paragraph("Service: ", style='List Bullet').add_run("report_service.py, pdf_generator.py").bold = True
    doc.add_paragraph(
        "Workflow: The AI's conclusions are mapped to concrete evidence (FDA counts, database matches). "
        "The backend formats this into a structured JSON. It also triggers pdf_generator.py to create downloadable PDF reports and Excel summaries."
    )

    # Section 5: API Endpoints Overview
    doc.add_heading('5. Core API Routers', level=1)
    doc.add_paragraph("auth.router (/auth): User authentication (Login/Register).", style='List Bullet')
    doc.add_paragraph("analyze.router (/analyze): Handling raw document uploads and segmentation.", style='List Bullet')
    doc.add_paragraph("report.router (/report): Triggering LLM reasoning and fetching final reports.", style='List Bullet')
    doc.add_paragraph("fda.router (/fda): Direct interfaces to fetch OpenFDA data charts.", style='List Bullet')
    doc.add_paragraph("chat.router (/chat): Conversational interface with the generated reports.", style='List Bullet')

    # Section 6: Local Development Setup
    doc.add_heading('6. Running the Project Locally', level=1)
    doc.add_paragraph("Backend:", style='List Bullet')
    p = doc.add_paragraph("cd backend\nsource venv/bin/activate\npython -m uvicorn main:app --reload")
    p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph("Milvus Lite:", style='List Bullet')
    p = doc.add_paragraph("./venv/bin/milvus-lite server --data-dir milvus/data")
    p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph("Frontend:", style='List Bullet')
    p = doc.add_paragraph("cd frontend_new\nnpm run dev")
    p.paragraph_format.left_indent = Inches(0.5)

    # Save
    doc.save('/Users/affu01/GRAD_PROJ_NEW/AI_Pharmacovigilance_KT_Document.docx')
    print("Documentation successfully generated.")

if __name__ == "__main__":
    create_document()
